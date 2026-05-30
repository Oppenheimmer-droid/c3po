from typing import Any, Dict, List
from typing import Dict
"""Document processing service for handling uploads and chunking."""

import os
import hashlib
import json
from typing import Optional, List, Tuple
from uuid import UUID, uuid4
from datetime import datetime
from pathlib import Path
import logging

import aiofiles
from sqlalchemy.ext.asyncio import AsyncSession
from pypdf import PdfReader
from docx import Document as DocxDocument

from app.models import Document, DocumentChunk
from app.repositories.documents import DocumentRepository, DocumentChunkRepository
from app.rag.vector_store import vector_store
from app.core.config import settings

logger = logging.getLogger(__name__)


class DocumentService:
    """Service for document upload, processing, and management."""
    
    UPLOAD_DIR = Path("uploads")
    
    def __init__(self, db: AsyncSession):
        self.db = db
        self.doc_repo = DocumentRepository(db)
        self.chunk_repo = DocumentChunkRepository(db)
        
        # Ensure upload directory exists
        self.UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
    
    def _get_tenant_upload_dir(self, tenant_id: UUID) -> Path:
        """Get tenant-specific upload directory."""
        tenant_dir = self.UPLOAD_DIR / str(tenant_id)
        tenant_dir.mkdir(parents=True, exist_ok=True)
        return tenant_dir
    
    def _generate_file_path(
        self, tenant_id: UUID, original_filename: str
    ) -> Tuple[str, str]:
        """Generate unique file path for uploaded document."""
        ext = Path(original_filename).suffix.lower()
        timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        hash_suffix = hashlib.md5(f"{original_filename}{timestamp}".encode()).hexdigest()[:8]
        filename = f"{timestamp}_{hash_suffix}{ext}"
        file_path = str(self._get_tenant_upload_dir(tenant_id) / filename)
        relative_path = f"uploads/{tenant_id}/{filename}"
        return file_path, relative_path
    
    async def save_upload(
        self, content: bytes, original_filename: str, tenant_id: UUID
    ) -> Tuple[str, str, str]:
        """Save uploaded file and return paths."""
        file_path, relative_path = self._generate_file_path(tenant_id, original_filename)
        
        mime_type = self._get_mime_type(original_filename)
        
        async with aiofiles.open(file_path, "wb") as f:
            await f.write(content)
        
        return file_path, relative_path, mime_type
    
    def _get_mime_type(self, filename: str) -> str:
        """Determine MIME type from filename."""
        ext = Path(filename).suffix.lower()
        mime_types = {
            ".pdf": "application/pdf",
            ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ".doc": "application/msword",
            ".txt": "text/plain",
            ".md": "text/markdown",
        }
        return mime_types.get(ext, "application/octet-stream")
    
    def _extract_text_pdf(self, file_path: str) -> List[Tuple[str, int]]:
        """Extract text from PDF with page numbers."""
        chunks = []
        reader = PdfReader(file_path)
        for page_num, page in enumerate(reader.pages, start=1):
            text = page.extract_text()
            if text:
                chunks.append((text, page_num))
        return chunks
    
    def _extract_text_docx(self, file_path: str) -> List[Tuple[str, None]]:
        """Extract text from DOCX."""
        doc = DocxDocument(file_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        return [("\n".join(full_text), None)] if full_text else []
    
    def _extract_text_txt(self, file_path: str) -> List[Tuple[str, None]]:
        """Extract text from plain text files."""
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
        return [(text, None)] if text.strip() else []
    
    def _chunk_text(
        self, text: str, page_number: Optional[int] = None
    ) -> List[Dict[str, Any]]:
        """Split text into semantic chunks."""
        chunks = []
        
        # Simple chunking by character count with overlap
        chunk_size = settings.CHUNK_SIZE
        overlap = settings.CHUNK_OVERLAP
        
        start = 0
        chunk_index = 0
        
        while start < len(text):
            end = start + chunk_size
            chunk_text = text[start:end]
            
            # Don't cut in the middle of a word
            if end < len(text) and chunk_text[-1] not in " \n\t.":
                last_space = chunk_text.rfind(" ")
                if last_space > chunk_size // 2:
                    chunk_text = chunk_text[:last_space]
                    end = start + last_space
            
            if chunk_text.strip():
                chunk_id = str(uuid4())
                chunks.append({
                    "id": chunk_id,
                    "content": chunk_text.strip(),
                    "chunk_index": chunk_index,
                    "page_number": page_number,
                    "start_char": start,
                    "end_char": end,
                })
                chunk_index += 1
            
            start = end - overlap if overlap > 0 else end
        
        return chunks
    
    async def create_document(
        self,
        tenant_id: UUID,
        user_id: UUID,
        filename: str,
        file_content: bytes,
        title: str,
        description: Optional[str] = None,
        subject_id: Optional[UUID] = None,
        topic_id: Optional[UUID] = None,
        difficulty: int = 1,
        tags: Optional[List[str]] = None,
    ) -> Document:
        """Create and save a new document."""
        # Save file
        file_path, relative_path, mime_type = await self.save_upload(
            file_content, filename, tenant_id
        )
        
        # Get page count
        page_count = None
        if mime_type == "application/pdf":
            try:
                reader = PdfReader(file_path)
                page_count = len(reader.pages)
            except Exception as e:
                logger.warning(f"Could not read PDF page count: {e}")
        
        # Create document record
        document = Document(
            tenant_id=tenant_id,
            uploaded_by=user_id,
            filename=relative_path,
            original_filename=filename,
            file_size=len(file_content),
            mime_type=mime_type,
            file_path=file_path,
            title=title,
            description=description,
            subject_id=subject_id,
            topic_id=topic_id,
            difficulty=difficulty,
            status="pending",
            page_count=page_count,
            tags=json.dumps(tags) if tags else None,
        )
        
        self.db.add(document)
        await self.db.flush()
        await self.db.refresh(document)
        
        return document
    
    async def process_document(
        self, document_id: UUID, tenant_id: UUID
    ) -> Tuple[int, List[str]]:
        """
        Process a document: extract text, chunk, and embed.
        
        Returns:
            Tuple of (chunk_count, vector_ids)
        """
        document = await self.doc_repo.get_by_id(document_id)
        if not document:
            raise ValueError(f"Document {document_id} not found")
        
        # Update status
        document.status = "processing"
        await self.db.flush()
        
        try:
            # Extract text based on file type
            text_chunks = []
            if document.mime_type == "application/pdf":
                text_chunks = self._extract_text_pdf(document.file_path)
            elif document.mime_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
                text_chunks = self._extract_text_docx(document.file_path)
            elif document.mime_type in ["text/plain", "text/markdown"]:
                text_chunks = self._extract_text_txt(document.file_path)
            
            if not text_chunks:
                raise ValueError("No text extracted from document")
            
            # Chunk all extracted text
            all_chunks = []
            for text, page_num in text_chunks:
                chunks = self._chunk_text(text, page_num)
                all_chunks.extend(chunks)
            
            # Store chunks in database
            db_chunks = []
            for chunk_data in all_chunks:
                chunk = DocumentChunk(
                    document_id=document_id,
                    tenant_id=tenant_id,
                    content=chunk_data["content"],
                    chunk_index=chunk_data["chunk_index"],
                    page_number=chunk_data["page_number"],
                    start_char=chunk_data["start_char"],
                    end_char=chunk_data["end_char"],
                )
                self.db.add(chunk)
                db_chunks.append(chunk)
            
            await self.db.flush()
            
            # Prepare chunks for vector store
            vector_chunks = []
            for db_chunk in db_chunks:
                vector_chunks.append({
                    "id": str(db_chunk.id),
                    "content": db_chunk.content,
                    "metadata": {
                        "tenant_id": str(tenant_id),
                        "document_id": str(document_id),
                        "subject_id": str(subject_id) if document.subject_id else None,
                        "topic_id": str(document.topic_id) if document.topic_id else None,
                        "page": db_chunk.page_number,
                        "chunk_index": db_chunk.chunk_index,
                    },
                })
            
            # Add to vector store
            vector_ids = await vector_store.add_chunks(
                tenant_id=str(tenant_id),
                chunks=vector_chunks,
            )
            
            # Update chunk records with vector IDs
            for i, db_chunk in enumerate(db_chunks):
                db_chunk.vector_id = vector_ids[i]
            
            # Update document status
            document.status = "completed"
            document.chunk_count = len(db_chunks)
            await self.db.flush()
            
            logger.info(f"Document {document_id} processed: {len(db_chunks)} chunks created")
            
            return len(db_chunks), vector_ids
            
        except Exception as e:
            document.status = "failed"
            document.error_message = str(e)
            await self.db.flush()
            logger.error(f"Error processing document {document_id}: {e}")
            raise
    
    async def delete_document(self, document_id: UUID, tenant_id: UUID) -> bool:
        """Delete a document and its associated chunks."""
        document = await self.doc_repo.get_by_id(document_id)
        if not document or document.tenant_id != tenant_id:
            return False
        
        # Delete from vector store
        await vector_store.delete_by_document(str(tenant_id), str(document_id))
        
        # Delete chunks from database
        await self.chunk_repo.delete_by_document(document_id)
        
        # Delete file
        if os.path.exists(document.file_path):
            os.remove(document.file_path)
        
        # Delete document record
        await self.doc_repo.delete(document_id)
        
        return True
    
    async def get_document(self, document_id: UUID, tenant_id: UUID) -> Optional[Document]:
        """Get a document by ID."""
        document = await self.doc_repo.get_by_id(document_id)
        if document and document.tenant_id == tenant_id:
            return document
        return None
    
    async def list_documents(
        self,
        tenant_id: UUID,
        skip: int = 0,
        limit: int = 20,
        status: Optional[str] = None,
    ) -> Tuple[List[Document], int]:
        """List documents for a tenant."""
        documents = await self.doc_repo.get_by_tenant(
            tenant_id, skip, limit, status
        )
        total = await self.doc_repo.count_by_tenant(tenant_id, status)
        return documents, total